import spacy
import docx
nlp = spacy.load("ru_core_news_sm")


def find_requisites(text):
    doc = nlp(text)

    customer_requisites = ""
    contractor_requisites = ""

    for ent in doc.ents:
        if "Заказчик" in ent.text:
            customer_requisites = ent.text
        elif "Подрядчик" in ent.text:
            contractor_requisites = ent.text

    return customer_requisites, contractor_requisites

def requisits_psrse():
    d = ''''''
    doc = docx.Document('C:\\Users\\Slava-Stat\\Desktop\\Проекты_Python\\Excel_parsing\\пример_файлов\\пример ворда (это геморойней)\\пример договоров\\пример ворда (это геморойней)\\пример договоров\\Договор Декор.docx')
    for paragraph in doc.paragraphs:
        d += paragraph.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                d += cell.text
    return d
# Пример текста документа договора
document_text = requisits_psrse()

customer_requisites, contractor_requisites = find_requisites(document_text)
print("Реквизиты заказчика:", customer_requisites)
print("Реквизиты подрядчика:", contractor_requisites)