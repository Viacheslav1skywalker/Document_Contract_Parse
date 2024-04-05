import re

import PyPDF2
import aspose.words as aw
from docx import Document
import os
import read_files
import write_in_file

class ParsingContractDocument:
    dict_values = {'название организации':None,'инн':None,
                   'кпп':None,'огрн':None,
                   'расчетный счет':None,'корреспондентский счет':None,
                   'бик':None,'адрес':None,
                   'телефон':None,'номер договора':None,
                   'дата договора':None,'кадастровый номер работ':None,
                   'виды работ':None,'стоимость работ':None,
                   'файл из которого взяты данные':None}
    def __init__(self,file_files):
        '''data - file or direcory name'''
        self.documnets_files = read_files.Open_files(file_files).documnets_files
        if not self.documnets_files:
            self.documnets_files = [file_files]


    def main_parse(self):
        all_values = []
        for i in self.documnets_files:
            print('текущий файл - ', i)
            self.dict_values['файл из которого взяты данные'] = i
            self.text_doc = self.method_defination(i)
            self.text_doc_lower = self.text_doc.lower()
            print()
            print('текущий файл - ',i)
            print(self.requisites_parse())
            print(self.kadastr_number())
            print(self.number_contract())
            print(self.working_cost())
            print(self.data_contract())
            print()
            print('result')
            print(self.dict_values)
            print('end')
            all_values.append(self.dict_values.copy())
            self.delete_data()
        return all_values

    def delete_data(self):
        for i in self.dict_values:
            self.dict_values[i] = None
    def method_defination(self,value):
        if value.endswith('.docx'):
            return self.extract_data_docx_file(value)
        elif value.endswith('.doc'):
            return self.extract_data_doc_file(value)
        elif value.endswith('.pdf'):
            return self.extract_data_pdf_file(value)
        else:
            raise FileNotFoundError(
                                    '''данный код может обрабатывать только файлы с 
                                        расширением docx, doc, pdf''')
    def transformation_data_replace_first_spaces(self,values:list):
        return [i.lstrip() for i in values]

    def extract_data_docx_file(self,path):
        '''Читаем docx файл'''
        shablons_start = [
            r'адреса.*сторон.*|подписи.*',
            r'рекв[еи]зиты.*?ст.*н?',
        ]
        doc = Document(path)
        full_text_from_tables = []
        full_text_from_document = []
        previous_text = ""
        # Извлечение текста из параграфов с указанием стиля
        for paragraph in doc.paragraphs:
            if paragraph.text:  # Проверяем, что текст параграфа не пустой
                full_text_from_document.append(paragraph.text)

        # Извлечение текста из таблиц
        for table in doc.tables:
            table_text = ''
            for row in table.rows:
                try:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text != previous_text:
                            table_text += cell_text + '\n'
                            full_text_from_tables.append(cell_text)
                            previous_text = cell_text
                except IndexError:
                    pass
        document_text = '\n'.join(full_text_from_document)
        all_res = []
        for sh in shablons_start:
            all_res += re.findall(sh, document_text.lower(), re.I | re.DOTALL)
        for text in all_res:
            document_text = document_text.lower().replace(text, '',4)
        dicts_values = {'tables_inf': self.search_requisites_text('\n'.join(full_text_from_tables)),
                        'text_inf': '\n'.join(full_text_from_document)}
        print(dicts_values['tables_inf'])
        print(dicts_values['text_inf'])

        # print('данные в таблицах')
        # print(dicts_values['text_inf'])
        # print(document_text + 'адреса и расчетные счета сторон' + '\n' + dicts_values['tables_inf'])
        return document_text + 'адреса и расчетные счета сторон' + '\n' + dicts_values['tables_inf']

    def search_requisites_text(self,text):
        search = re.findall(r'заказчик.*?инн.*', text, re.I | re.DOTALL)
        if search:
            text = search[0].lower()
            if 'подрядчик' in  text[:text.index('инн')]:
                text = text.replace('подрядчик','')
            return text

    def extract_data_pdf_file(self, file_path):
        text = ""
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
        return text
    def extract_data_doc_file(self,doc_path):
        docx_path = 'result_information_in_docx_file.docx'
        if os.path.exists(docx_path):
            os.remove(docx_path)
        # Загружаем документ .doc
        try:
            doc = aw.Document(doc_path)
        except:
            print('ошибка доступа к файлу, файл занят другим процессом')
            raise RuntimeError('закройте файл чтобы продолжить')
        # Сохраняем документ в формате .docx
        doc.save(docx_path)
        print(f'Файл {doc_path} успешно преобразован в файл {docx_path}')

        # Чтение содержимого нового файла .docx
        return self.extract_data_docx_file(docx_path)
    def parsing_name_customer(self,text):

        shablons = [r'заказчик:.+?(.+?".+?")',r'заказчик:(.+?»)',r'(ООО.+«.+»|Общество с ограниченной ответственностью.+«.+»).+именуемое в дальнейшем "заказчик"'
                    ]
        all_res = []
        for shablon in shablons:
            all_res += re.findall(shablon,text,re.I | re.DOTALL)
        if not all_res:
            print('Имя заказчика не найдено')
        print('имя заказчика:')
        self.dict_values['название организации'] = self.return_smallest_index_value(text,self.transformation_data_replace_first_spaces(all_res))


    def kadastr_number(self):
        """парсинг номера кадастровых работ"""
        shablons = [r'КН\s*\d+:\d+:\d+:\d+']
        res_lst = []
        for shablon in shablons:
            res_lst += re.findall(shablon,self.text_doc,re.I|re.DOTALL)
        print('кадастровый номер:')
        if res_lst:
            self.dict_values['кадастровый номер работ'] = res_lst[0]

    def number_contract(self):
        shablons = [r'договор\s*№\s*.+']
        res_lst = []
        for shablon in shablons:
            res_lst  += re.findall(shablon,self.text_doc,re.I)
        print('номер договора:')
        print(res_lst[0])
        if res_lst:
            self.dict_values['номер договора'] = res_lst[0]


    def working_cost(self):
        shablons = [r'\d+\s*\d+,?\d*\s*?\(.*?\)\s*руб',r'\d+\s*\d*,?\d*\s*руб']
        res_lst = []
        text = self.element_that_text_startet([r'цена договора'])
        for shablon in shablons:
            res_lst += re.findall(shablon,text,re.I|re.DOTALL)
        print('цена договора')
        if res_lst:
            self.dict_values['стоимость работ'] = self.return_smallest_index_value(text,res_lst)
    def data_contract(self):
        shablons = [r'[«"“]?\s*?\d{1,2}\s*?["»”]?\s+\w+\s*\d+\s*?г',r'\d{1,2}\s+\w+\s+\d{4}\s+г\.']
        res_lst = []
        for shablon in shablons:
            res_lst += re.findall(shablon, self.text_doc, re.I | re.DOTALL)
        print("дата договора")
        if res_lst:
            self.dict_values['дата договора'] = res_lst[0]


    def addres(self,text):
        # Шаблон для почтового адреса
        shablons = [
            r'заказчик.*?(юр.*?ий\s*адрес.*?)тел',
            r'заказчик.*?(юр.*?ий\s*адрес.*?)инн',
            r'заказчик.*?(адрес.+?)инн',
            r'заказчик.*(поч.*ый\s*адрес.+?)инн',
            r'заказчик.*?(юр.*?\s*адрес.*?)огрн',
            r'заказчик.*?(юр.*?\s*адрес.*?)р/с'
                ]
        lst_res = []
        for shablon in shablons:
            lst_res += re.findall(shablon, text, re.I | re.DOTALL)
        print('результат парсинга адресса')
        if not lst_res:
            print('Адрес не найден')
        self.dict_values['адрес'] = self.return_smallest_index_value(text,lst_res)

    def inn_parse(self,text):
        inn_shablons = [r'инн.*?(\d{10})\D',r'инн.*?(\d{12})\D',
                        r'инн.*?\D(\d{10})\D']
        all_search_values = []
        for shablon in inn_shablons:
            all_search_values += re.findall(shablon,text,re.I|re.DOTALL)
        print('результат парсинга инн:')
        print(self.return_smallest_index_value(text,all_search_values))
        if not all_search_values:
            print('ИНН не найден')
        else:
            self.dict_values['инн'] = self.return_smallest_index_value(text,all_search_values)
            print(self.dict_values)
    def bic_parse(self,text):
        bic_shablons = [r'бик.*?\D(\d{9})\D']
        all_search_values = []
        for shablon in bic_shablons:
            all_search_values += re.findall(shablon, text,re.I|re.DOTALL)
        print('результат парсинга бик:')
        if not all_search_values:
            print('БИК не найден')
        self.dict_values['бик'] = self.return_smallest_index_value(text, all_search_values)

    def kpp_parse(self,text):
        kpp_shablons = [r'кпп.*?\D'
                        r'(\d{9})\D']
        all_search_values = []
        for shablon in kpp_shablons:
            all_search_values += re.findall(shablon, text,re.I|re.DOTALL)
        if not all_search_values:
            print('КПП не найден')
        print('результат парсинга кпп:')
        self.dict_values['кпп'] = self.return_smallest_index_value(text, all_search_values)

    def ogrn_parse(self,text):
        ogrn_shablons = [r'огрн.*?\D(\d+?)\D']
        all_search_values = []
        for shablon in ogrn_shablons:
            all_search_values += re.findall(shablon, text,re.I|re.DOTALL)
        if not all_search_values:
            print('огрн не найден')
        print('результат парсинга огрн:')
        self.dict_values['огрн'] = self.return_smallest_index_value(text, all_search_values)

    def payment_account(self,text):
        r_s_shablons = [r'р[а-я]*/с[а-я]*.*?\D(\d{20}\d*)\D',r'рас.*?сч.*?\D(\d{20}\d*)\D']
        all_search_values = []
        for shablon in r_s_shablons:
            all_search_values += re.findall(shablon, text,re.I|re.DOTALL)
        if not all_search_values:
            print('расчетный счет не найден')
        print(text)
        print('результат парсинга расчетного счета:')
        print(all_search_values)
        self.dict_values['расчетный счет'] = self.return_smallest_index_value(text, all_search_values)

    def correspondent_account(self,text):
        k_a_shablons = [r'к[а-я]*/с[а-я]*.*?\D(\d{20})\D',r'кор.*?сч.*?\D(\d{20})\D']
        all_search_values = []
        for shablon in k_a_shablons:
            all_search_values += re.findall(shablon, text,re.I|re.DOTALL)
        if not all_search_values:
            print('корреспондентский счет не найден')
            return

        print('результат парсинга корреспондентского счета:')
        self.dict_values['корреспондентский счет'] = self.return_smallest_index_value(text, all_search_values)

    def phone_number_parse(self,text):
        end_phone_number = ['подрядчик']
        string_with_telephon_data = r'тел\.|телефон[:]|тел'
        search = re.findall(string_with_telephon_data,text,re.I|re.DOTALL)
        if not search:
            print('Выражение с таким шаблоном поиска телефона не найдено')
        num_phone = ''
        symb = r'\d|\s|-|\(|\)'
        if search:
            index_invalid_simbs = ''
            for i in text[text.index(search[0])+len(search[0]):]:
                serch = re.findall(symb,i,re.I|re.DOTALL)
                if serch:
                    num_phone += i
                else:
                    if num_phone:
                        break
                    if index_invalid_simbs == '':
                        index_invalid_simbs += i
                        continue
                    elif not re.findall(symb,index_invalid_simbs,re.I|re.DOTALL):
                        continue
                    break
            print('текст поиска телефона')
            if not num_phone:
                print('Номер телефона не найден')
        self.dict_values['телефон'] = num_phone

    def requisites_parse(self):
        values_requisites = ['инн','телефон','тел.','огрн',
                             'р/с','рас','расчетный','почтовый адрес',
                             'юридический адрес','кор/с',"к/с"]
        end_notify_word = ['подрядчик']
        shablons_start = [
            r'адреса.*?сторон',
            r'рекв[еи]зиты.*?сторон',
        ]
        text = self.search_certain_start_end(self.text_doc,shablons_start,end_notify_word)
        if not text:
            print('Стартовое значене с такими шаблонами не найдено')
        print(self.inn_parse(text))
        print(self.kpp_parse(text))
        print(self.ogrn_parse(text))
        print(self.parsing_name_customer(text))
        print(self.addres(text))
        print(self.correspondent_account(text))
        print(self.payment_account(text))
        print(self.bic_parse(text))
        print(self.phone_number_parse(text))




    def check_simillar_index_start(self,lst_index_start:list,min_value):
        '''проверяем есть ли в элементе одинаковые значения индекса старта текста'''
        if lst_index_start.count(min_value) > 1:
            return True

    def func_search(self,shablons_value,shablons_name_start=None):
        """главная функция: шаблон поиска значений в договоре"""
        start_word = None
        for i in shablons_name_start:
            a = re.findall(i,self.text_doc,re.I|re.DOTALL)
            if a:
                start_word = a[0]
        if start_word:
            text = self.text_doc.lower()[self.text_doc.lower().index(start_word):]
            for i in shablons_value:
                a = re.findall(i,text)
                if a:
                    pass
                # функция не доделана



    def find_smallest_length(self,values:list,texts:list):
        index_min = min(values)
        lsts_comparison = []
        for i in range(len(texts)):
            if values[i] == index_min:
                lsts_comparison.append(texts[i])
        len_check = None
        for i in lsts_comparison:
            if i == 0:
                len_check = texts[i]
            if len(texts[i]) < len(len_check):
                len_check = texts[i]
        return len_check

    def working_type(self):
        pass

    def element_that_text_startet(self,shablons):
        for shablon in shablons:
            search = re.findall(shablon,self.text_doc_lower,re.I | re.DOTALL)
            if search:
                return self.text_doc_lower[self.text_doc_lower.index(search[0]):]
        print('шаблон соответствующий выражению в тексте не найден')
        # в противном случае возвращаем тот же исходный текст
        return self.text_doc_lower



    def search_certain_start_end(self,text,start_element,ended_element):
        ind_start = None
        ind_end = None
        for i in start_element:
            search = re.findall(i,text,re.I|re.DOTALL)
            if search:
                ind_start = search[0]
                break
        text_search_now = text[text.index(ind_start):]
        for i in ended_element:
            search = re.findall(i, text_search_now, re.I | re.DOTALL)
            if search:
                ind_end = search[0]
                break
        if ind_end:
            return text_search_now[:text_search_now.index(ind_end)].lower()
        return text_search_now



    def return_smallest_index_value(self,text:str,values:list):
        '''возвращает значение которое находится ближе всех по индексу в тексте и
        которое самое маленькое по размеру'''

        least_index = None
        for i in range(len(values)):
            if i == 0:
                least_index = values[i]
                continue
            if values[i] == '' or values[i].isspace():
                continue
            if text.index(values[i]) < text.index(least_index):
                least_index = values[i]
                continue
            elif text.index(values[i]) == text.index(least_index) and least_index != None:
                least_index = values[i] if len(values[i]) <= len(least_index) else least_index
        return least_index


    def test(self):
        a = re.findall(r'юр.+ий\s*адрес',self.text_doc_lower,re.I|re.DOTALL)
        text = self.text_doc.lower()[self.text_doc.lower().index(a[0]):]
        print('найденный элемент:')
        print(a[0])
        if a:
            print('результат')
            print(self.text_doc.lower()[self.text_doc.lower().index(a[0]):self.text_doc.lower().index('тел')])




#
# d = ParsingContractDocument('C:\\Users\Slava-Stat\Desktop\Проекты_Python\Excel_parsing\пример_файлов\пример ворда (это геморойней)\пример договоров\пример ворда (это геморойней)\пример договоров\Договор схема.docx')
# # print(d.text_doc)
# # print(d.addres())
# print(d.main_parse())

# obj = ParsingContractDocument()
# obj.main_parse()

