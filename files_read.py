from docx import Document

class ReadFiles:


    def file_parse(self):
        pass

    def extract_text_from_docx(self,file_path):
        doc = Document(file_path)
        full_text = []
        previous_text = ""

        # Извлечение текста из параграфов с указанием стиля
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para.style.name and para_text != previous_text:
                full_text.append(f"{para.style.name}: {para_text}")
                previous_text = para_text
            elif para_text != previous_text:
                full_text.append(para_text)
                previous_text = para_text

        # Извлечение текста из таблиц
        table_texts = set()  # Для хранения уникальных текстов из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    table_texts.add(cell_text)

        full_text.extend(table_texts)

        return '\n'.join(full_text)