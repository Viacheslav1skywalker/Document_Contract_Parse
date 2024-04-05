
import aspose.words as aw
from docx import Document
import os
class Open_files:
    def __init__(self,path):
        all_finding_files = self.add_files(path)
        self.excel_files = all_finding_files['excel_files']
        self.documnets_files = all_finding_files['documents_files']


    def add_files(self,path):
        '''пробегаемся по всей файловой системе начиная от конкретной папки
            и возвращаем список всех путей к файлу которые имеют слово "договор" в своем названии'''
        found_files_excel = []
        found_files_documents = []
        # Рекурсивно обходим все файлы в заданной папке и ее подпапках
        for foldername, subfolders, filenames in os.walk(path):

            for filename in filenames:
                if 'дог' in filename.lower():
                    if filename.endswith('.xlsx') or filename.endswith('.xlsm'):
                        found_files_excel.append(os.path.join(foldername,filename))
                    elif filename.endswith('.doc') or filename.endswith('.docx'):
                        found_files_documents.append(os.path.join(foldername,filename))
        return {'excel_files':found_files_excel,'documents_files':found_files_documents}

    def file_or_files(self, file):
        if os.path.isfile(file):
            return file
        elif os.path.isdir(file):
            return self.add_files(file)
        else:
            print('файл не найден или путь к файлу указан неверно')
            raise FileNotFoundError('Ссылка на файл не найдена')

    def read_doc_files(self,doc_path):
        docx_path = 'result_information_in_docx_file.docx'
        if os.path.exists(docx_path):
            os.remove(docx_path)
        # Загружаем документ .doc
        doc = aw.Document(doc_path)
        # Сохраняем документ в формате .docx
        doc.save(docx_path)

        print(f'Файл {doc_path} успешно преобразован в файл {docx_path}')

        # Чтение содержимого нового файла .docx
        docx = aw.Document(docx_path)
        for paragraph in docx.get_child_nodes(aw.NodeType.PARAGRAPH, True):
            print(paragraph.get_text().strip())

        doc = Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return '\n'.join(full_text)


d = Open_files(r'C:\\Users\\Slava-Stat\\Desktop\\договора\\ДОГОВОР 2023-0030 сломан.xlsm')
print(d.excel_files)













